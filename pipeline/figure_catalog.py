# figure_catalog.py
# Build a figure catalog from DOCX files using python-docx.
# Extracts both caption text AND embedded images directly from the DOCX,
# eliminating asset matching issues.

from __future__ import annotations
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Hebrew caption patterns: "איור X.Y:" or "טבלה X.Y:"
_CAPTION_RE = re.compile(
    r'(איור|טבלה)\s+(\d+)\.(\d+)([a-zA-Z]*)\s*[:：－–—\-]\s*(.*)',
    re.UNICODE,
)

# Content type to extension mapping
_CONTENT_TYPE_EXT = {
    'image/jpeg': '.jpg',
    'image/png': '.png',
    'image/webp': '.webp',
    'image/svg+xml': '.svg',
    'image/gif': '.gif',
    'image/tiff': '.tiff',
    'image/bmp': '.bmp',
    'image/x-emf': '.emf',
    'image/x-wmf': '.wmf',
}


@dataclass
class FigureInfo:
    key: str               # e.g. "7.1" or "7.1a"
    kind: str              # "איור" or "טבלה"
    chapter: int
    number: int
    suffix: str            # e.g. "a", "b", or ""
    caption: str           # full caption text after the colon
    docx_para_index: int   # paragraph index in DOCX for ordering
    image_blob: Optional[bytes] = None   # embedded image data
    image_ext: str = ".jpg"              # image file extension


def _extract_image_from_para(doc, para) -> tuple[Optional[bytes], str]:
    """Extract the first embedded image from a paragraph.

    Returns (blob, extension) or (None, '') if no image found.
    """
    for run in para.runs:
        drawings = run._element.findall(qn('w:drawing'))
        for drawing in drawings:
            blips = drawing.findall('.//' + qn('a:blip'))
            for blip in blips:
                embed = blip.get(qn('r:embed'))
                if embed:
                    rel = doc.part.rels.get(embed)
                    if rel and hasattr(rel, 'target_part'):
                        tp = rel.target_part
                        ext = _CONTENT_TYPE_EXT.get(tp.content_type, '.png')
                        return tp.blob, ext
    return None, ''


def _find_image_near_caption(doc, caption_idx: int, max_distance: int = 5) -> tuple[Optional[bytes], str]:
    """Search paragraphs near a caption for an embedded image.

    Searches below first (most common), then above.
    """
    paras = doc.paragraphs

    # Search below (image usually follows caption)
    for j in range(caption_idx + 1, min(len(paras), caption_idx + 1 + max_distance)):
        blob, ext = _extract_image_from_para(doc, paras[j])
        if blob:
            return blob, ext

    # Search above (sometimes image comes before caption)
    for j in range(max(0, caption_idx - max_distance), caption_idx):
        blob, ext = _extract_image_from_para(doc, paras[j])
        if blob:
            return blob, ext

    return None, ''


def build_catalog_from_docx(docx_path: Path) -> dict[str, FigureInfo]:
    """Extract figure/table captions and their embedded images from a DOCX file.

    Returns a dict mapping figure keys like "7.1" to FigureInfo objects.
    """
    if not HAS_DOCX:
        return {}

    doc = Document(str(docx_path))
    catalog: dict[str, FigureInfo] = {}

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        m = _CAPTION_RE.match(text)
        if not m:
            continue

        kind, ch, num, suffix, rest = m.groups()
        key = f"{int(ch)}.{int(num)}{suffix or ''}"

        # Collect continuation: sometimes caption spans multiple paragraphs
        caption_parts = [rest.strip()]
        for j in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
            next_para = doc.paragraphs[j]
            next_text = next_para.text.strip()
            if not next_text:
                break
            if _CAPTION_RE.match(next_text):
                break
            style = next_para.style
            if style and style.name and 'heading' in style.name.lower():
                break
            # Don't include source/credit lines in caption continuation
            # if they are image-bearing paragraphs
            blob, _ = _extract_image_from_para(doc, next_para)
            if blob:
                break
            caption_parts.append(next_text)

        # Find the embedded image near this caption
        image_blob, image_ext = _find_image_near_caption(doc, idx)

        catalog[key] = FigureInfo(
            key=key,
            kind=kind,
            chapter=int(ch),
            number=int(num),
            suffix=suffix or "",
            caption=" ".join(caption_parts),
            docx_para_index=idx,
            image_blob=image_blob,
            image_ext=image_ext,
        )

    return catalog


def build_catalog_for_chapter(
    raw_dir: Path, chapter_num: int
) -> dict[str, FigureInfo]:
    """Build catalog for a specific chapter by finding its DOCX file."""
    for p in raw_dir.glob("*.docx"):
        stem = p.stem.lower().replace(" ", "")
        m = re.search(r'ch(\d+)', stem)
        if m and int(m.group(1)) == chapter_num:
            return build_catalog_from_docx(p)
    return {}


def save_catalog_images(
    catalog: dict[str, FigureInfo],
    media_dir: Path,
) -> dict[str, Path]:
    """Save all extracted images to the media directory.

    Returns a dict mapping figure keys to their saved file paths.
    """
    media_dir.mkdir(parents=True, exist_ok=True)
    saved: dict[str, Path] = {}

    for key, info in catalog.items():
        if info.image_blob is None:
            continue

        # Canonical filename: ch_num.ext (e.g. 7_1.jpg)
        out_name = f"{info.chapter}_{info.number}{info.suffix}{info.image_ext}"
        target = media_dir / out_name

        target.write_bytes(info.image_blob)
        saved[key] = target

    return saved
