# docx_processing.py
# Utilities that pre-process .docx files before Pandoc conversion.
#
# This module contains the "convert_numbered_to_headings_keep_bullets" logic
# that promotes *numbered* list items to Heading styles, while keeping bullets
# as list items. It also avoids touching paragraphs with drawings/captions.
#
# Extracted/refactored from convert_headers_from_numbering.py (pipeline-original).

import re
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from lxml import etree

# Namespaces
W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
V_NS   = "urn:schemas-microsoft-com:vml"
WP_NS  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
A_NS   = "http://schemas.openxmlformats.org/drawingml/2006/main"

NS = {'w': W_NS}

CAPTION_RE = re.compile(
    r'^\s*(טבלה|איור|תמונה|Fig\.?|Figure|Illustration)\b.*$', re.IGNORECASE
)

def build_numbering_maps(doc):
    """Parse numbering.xml and return (num_to_abs, level_info) maps.
    level_info[(absId, ilvl)] = (numFmt, lvlText, sym)
    """
    numpart = doc.part.numbering_part
    if numpart is None:
        return {}, {}
    root = etree.fromstring(numpart.blob)
    num_to_abs = {}
    for num in root.xpath('.//w:num', namespaces=NS):
        numId = (num.xpath('./@w:numId', namespaces=NS) or [None])[0]
        absId = (num.xpath('./w:abstractNumId/@w:val', namespaces=NS) or [None])[0]
        if numId:
            num_to_abs[numId] = absId
    level_info = {}
    for absNum in root.xpath('.//w:abstractNum', namespaces=NS):
        absId = (absNum.xpath('./@w:abstractNumId', namespaces=NS) or [None])[0]
        for lvl in absNum.xpath('./w:lvl', namespaces=NS):
            ilvl    = (lvl.xpath('./@w:ilvl', namespaces=NS) or [None])[0]
            numFmt  = (lvl.xpath('./w:numFmt/@w:val', namespaces=NS) or [''])[0]
            lvlText = (lvl.xpath('./w:lvlText/@w:val', namespaces=NS) or [''])[0]
            sym     = (lvl.xpath('./w:sym/@w:val', namespaces=NS) or [''])[0]
            if absId is not None and ilvl is not None:
                level_info[(absId, ilvl)] = (numFmt, lvlText, sym)
    return num_to_abs, level_info

def is_bullet(num_to_abs, level_info, numId_val, ilvl_val):
    """Decide bullet vs numbered for a (numId, ilvl) level."""
    absId = num_to_abs.get(numId_val)
    if absId is None:
        # Unknown numbering; default to 'numbered' so section titles don't get missed.
        return False
    numFmt, lvlText, sym = level_info.get((absId, ilvl_val), ('','',''))
    if (numFmt or '').lower() == 'bullet':
        return True
    if lvlText and ('%' not in lvlText):  # glyph bullet (no number slot)
        return True
    if sym:  # explicit glyph
        return True
    return False

def get_or_make_heading_styles(doc):
    """Return a dict {level:int -> style} for Heading 1..9, and the Normal style."""
    styles = {s.name: s for s in doc.styles}
    heading_styles = {}
    for i in range(1, 10):
        name = f"Heading {i}"
        if name in styles and styles[name].type == WD_STYLE_TYPE.PARAGRAPH:
            h = styles[name]
        else:
            # Create if missing (rare)
            h = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        heading_styles[i] = h
        # Make headings visually non-bold in the DOCX (cosmetic; Pandoc will output plain #)
        try:
            heading_styles[i].font.bold = False
        except Exception:
            pass
    normal_style = styles.get("Normal", None)
    if normal_style is None:
        normal_style = doc.styles.add_style("Normal", WD_STYLE_TYPE.PARAGRAPH)
    return heading_styles, normal_style

def strip_numpr(p):
    """Remove w:numPr from a paragraph if present (detaches numbering)."""
    p_elm = p._p
    pPr = p_elm.pPr if p_elm.pPr is not None else p_elm.get_or_add_pPr()
    numPr = pPr.find(qn('w:numPr'))
    if numPr is not None:
        pPr.remove(numPr)
        return True
    return False

def strip_para_direct_char_props(p):
    """
    Remove paragraph-level direct character props (<w:pPr><w:rPr>) that might enforce bold/italic,
    using safe .find() lookups.
    """
    p_elm = p._p
    pPr = p_elm.pPr if p_elm.pPr is not None else None
    if pPr is None:
        return
    rPr = pPr.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in ('w:b','w:bCs','w:i','w:iCs','w:u','w:strike','w:dstrike',
                'w:smallCaps','w:caps','w:vertAlign','w:highlight','w:color',
                'w:sz','w:szCs','w:rStyle','w:em'):
        child = rPr.find(qn(tag))
        if child is not None:
            rPr.remove(child)

def strip_run_inline_formatting(run):
    """Clear bold/italic and remove low-level rPr flags (so Pandoc won't emit **/*** in headings)."""
    try:
        run.bold = False
        run.italic = False
    except Exception:
        pass
    try:
        if run.style is not None:
            name = getattr(run.style, "name", None)
            if name in ("Strong", "Emphasis", "Intense Emphasis", "Intense Quote", "Quote"):
                run.style = None
    except Exception:
        pass
    rPr = getattr(run._r, "rPr", None)
    if rPr is not None:
        for tag in ('w:b','w:bCs','w:i','w:iCs','w:u','w:strike','w:dstrike',
                    'w:smallCaps','w:caps','w:vertAlign','w:highlight','w:color',
                    'w:sz','w:szCs','w:rStyle','w:em'):
            child = rPr.find(qn(tag))
            if child is not None:
                rPr.remove(child)

def strip_all_inline_formatting_in_heading(p):
    """Nuke any inline emphasis inside a heading so Pandoc won't emit *, **, or *** on heading lines."""
    strip_para_direct_char_props(p)
    for r in p.runs:
        strip_run_inline_formatting(r)

def para_has_drawing(p):
    """
    Return True if the paragraph contains any drawing/picture (w:drawing, w:pict, v:shape, etc.).
    Avoids xpath namespaces kwarg (python-docx wraps lxml); use tag checks instead.
    """
    el = p._p
    for node in el.iter():
        tag = node.tag
        if tag in (
            f"{{{W_NS}}}drawing",
            f"{{{W_NS}}}pict",
            f"{{{W_NS}}}object",
            f"{{{WP_NS}}}inline",
            f"{{{PIC_NS}}}pic",
            f"{{{V_NS}}}shape",
            f"{{{V_NS}}}group",
        ):
            return True
    return False

def is_captionish(p):
    """Heuristic: text looks like a figure/table caption OR style is caption."""
    txt = (p.text or "").strip()
    if CAPTION_RE.match(txt):
        return True
    sty = getattr(p.style, "name", "") or ""
    if sty and "caption" in sty.lower():
        return True
    return False

def convert_numbered_to_headings_keep_bullets(input_docx, output_docx=None):
    """
    If output_docx is None, overwrites input_docx in place (for pipelines that pass a single path).
    """
    out_path = output_docx or input_docx

    doc = Document(input_docx)
    num_to_abs, level_info = build_numbering_maps(doc)
    heading_styles, normal_style = get_or_make_heading_styles(doc)

    # Remove stray paras that are just "\\" (cleanup from Pandoc sometimes)
    for p in list(doc.paragraphs):
        if p.text.strip() == "\\":
            el = p._element
            el.getparent().remove(el)

    for p in doc.paragraphs:
        # Already a Heading → strip numbering + inline emphasis
        if p.style.name.startswith('Heading'):
            strip_numpr(p)
            strip_all_inline_formatting_in_heading(p)
            continue

        # Paragraphs with drawings/pictures must never become headings
        if para_has_drawing(p):
            strip_numpr(p)       # detach any accidental numbering binding
            p.style = normal_style
            continue

        # Caption-like lines must remain normal (not headings)
        if is_captionish(p):
            strip_numpr(p)
            p.style = normal_style
            continue

        # Check numbering for non-heading paragraphs
        p_elm = p._p
        pPr = p_elm.pPr if p_elm.pPr is not None else p_elm.get_or_add_pPr()
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue

        numId = numPr.find(qn('w:numId'))
        ilvl  = numPr.find(qn('w:ilvl'))
        numId_val = numId.get(qn('w:val')) if numId is not None else None
        ilvl_val  = ilvl.get(qn('w:val')) if ilvl is not None else None
        if ilvl_val is None or numId_val is None:
            continue

        # Bullet or numbered?
        if is_bullet(num_to_abs, level_info, numId_val, ilvl_val):
            # It's a bullet/list level → leave as-is.
            continue

        # Numbered level → convert to Heading {ilvl+1}, strip inline emphasis.
        level = min(int(ilvl_val) + 1, 9)
        strip_numpr(p)  # detach numbering binding

        # If there's no visible text (pure image, blank), keep as Normal
        if not (p.text or "").strip():
            p.style = normal_style
            continue

        # If this numbered paragraph is actually a caption-ish line, keep it Normal
        if is_captionish(p):
            p.style = normal_style
            continue

        # Convert to heading and strip inline emphasis
        p.style = heading_styles[level]
        strip_all_inline_formatting_in_heading(p)

    doc.save(out_path)
