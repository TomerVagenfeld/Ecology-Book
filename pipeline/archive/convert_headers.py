import sys
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

WNS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

def _get_numfmt(doc, numId_val, ilvl_val):
    """Return numFmt (e.g., 'bullet', 'decimal', 'upperRoman', ...) for (numId, ilvl)."""
    numbering = doc.part.numbering_part.element
    # find <w:num> with this numId
    nums = numbering.xpath(f'.//w:num[w:numId[@w:val="{numId_val}"]]', namespaces=WNS)
    if not nums:
        return None
    abstractNumId = nums[0].xpath('./w:abstractNumId/@w:val', namespaces=WNS)
    if not abstractNumId:
        return None
    absId = abstractNumId[0]
    # find abstractNum + level
    lvls = numbering.xpath(
        f'.//w:abstractNum[@w:abstractNumId="{absId}"]/w:lvl[@w:ilvl="{ilvl_val}"]/w:numFmt/@w:val',
        namespaces=WNS
    )
    return lvls[0] if lvls else None
# TODO: convert headers to numbers (preserve original numbering?)
def convert_list_paragraphs_to_headings(input_docx, output_docx=""):
    doc = Document(input_docx)
    if output_docx == "":
        output_docx = input_docx
    styles = doc.styles

    # 1) Remove any paragraph that is just a backslash
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip() == '\\':
            # Remove this paragraph element from the document
            p = paragraph._element
            p.getparent().remove(p)

    # Reference to Normal style
    normal_style = styles['Normal']

    # Ensure Heading 1–9 styles exist (or create them)
    heading_styles = {}
    for i in range(1, 10):
        name = f'Heading {i}'
        if name in styles:
            heading_styles[i] = styles[name]
        else:
            heading = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            heading.base_style = styles['Heading 1']
            heading_styles[i] = heading

    # Process remaining paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.style.name != 'List Paragraph':
            continue

        # Build or get paragraph properties
        p_elm = paragraph._p
        pPr = p_elm.pPr if p_elm.pPr is not None else p_elm.get_or_add_pPr()

        # Find numbering properties
        numPr = pPr.find(qn('w:numPr'))
        if numPr is None:
            continue

        # Determine list level (if present)
        ilvl = numPr.find(qn('w:ilvl'))
        if ilvl is None or ilvl.get(qn('w:val')) is None:
            # No valid level: strip numbering and reset to Normal
            pPr.remove(numPr)
            paragraph.style = normal_style
            continue

        level = int(ilvl.get(qn('w:val'))) + 1  # Word levels are zero-based

        # Remove numbering so Pandoc won't emit list markers
        pPr.remove(numPr)

        # Skip blank text (avoid empty headings)
        text = paragraph.text.strip()
        if not text:
            paragraph.style = normal_style
            continue

        # Apply heading style for this level
        style = heading_styles.get(level)
        if style:
            paragraph.style = style
        else:
            print(f"Warning: no Heading style for level {level}")

    doc.save(output_docx)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_headers.py input.docx output.docx")
    else:
        convert_list_paragraphs_to_headings(sys.argv[1], sys.argv[2])
